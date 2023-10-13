import streamlit as st
import requests
import openai
from pdfreader import SimplePDFViewer
from fpdf import FPDF
from docx import Document
from bs4 import BeautifulSoup
from io import BytesIO


st.title('AI Cover Letter Generator')

st.image("bot.png", use_column_width=True)

# Function to extract text from PDF
def extract_text_from_pdf(file):
    viewer = SimplePDFViewer(file)
    text = ""
    for canvas in viewer:
        text += " ".join(canvas.strings)
    return text

# Function to extract text from Word document
def extract_text_from_docx(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# Function to extract text from URL
def extract_text_from_url(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    return soup.get_text()

# Function to generate cover letter using OpenAI model
def generate_cover_letter(api_key, cv, job_description):
    openai.api_key = api_key
    response = openai.ChatCompletion.create(
    model=model_choice,
    messages=[
    {
      "role": "system",
      "content": """Write a Cover Letter for the role described in “Job Description” that is personalised to the details in contained within the attached CV.
      
      You cover letter should be written in the following formatting:

      **Date of Letter**
      **Contact Name**  
      **Contact Title**  
      **Company Name**  
      **Street Address**
      **City, State, Zip Code**
      
      **Dear : _name of hiring manager or recruiter_**
      
      [//]: # (Opening Paragraph)
      Clearly state why you're writing, name the position or type of work you're exploring and, where applicable, how you heard about the position or organization. A summary statement may work well here by including three reasons you think you would be a good fit for the opportunity.
      
      [//]: # (Middle Paragraph(s))
      Explain why you are interested in this employer and your reasons for desiring this type of work. If you've had relevant school or work experience, be sure to point it out with one or two key examples; but do not reiterate your entire resume. Emphasize skills or abilities that relate to the job. Be sure to do this in a confident manner and remember that the reader will view your letter as an example of your writing skills.
      
      [//]: # (closing Paragraph(s))
      Reiterate your interest in the position, and your enthusiasm for using your skills to contribute to the work of the organization. Thank the reader for their consideration of your application, and end by stating that you look forward to the opportunity to further discuss the position.
      

      Sincerely
      
      _Your name typed_
      """
    },
    {
      "role": "user",
      "content": f"Generate a cover letter based on the CV:\n{cv}\n\nAnd the job description:\n{job_description}\n\nCover Letter:"
    }
    ],
    temperature=1,
    max_tokens=4000,
    )
    return response.choices[0]['message']['content']

# For PDF create
def generate_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size = 12)
    pdf.multi_cell(0, 10, text)
    pdf_output = pdf.output(dest="S")
    return bytes(pdf_output)

# For DOCX create
def generate_docx(text):
    doc = Document()
    doc.add_paragraph(text)
    docx_output = BytesIO()
    doc.save(docx_output)
    docx_output.seek(0)
    return bytes(docx_output.read())

with st.sidebar:

    st.title('Settings')

    api_key = st.text_input('Enter your OpenAI API Key', type="password")
    
    model_choice = st.selectbox('Choose the GPT model:', ['gpt-4', 'gpt-3.5-turbo-16k'])

    st.write('This bot uses the power of Python and OpenAI to generate smart cover letters catered to your CV and the company\s the job advertisment.')

# Dropdown to choose whether to display extracted text from CV
display_cv_text = st.selectbox('Display extracted text from CV?', ['Yes', 'No'])

# Dropdown to choose whether to display extracted text from URL
display_url_text = st.selectbox('Display extracted text from job description URL?', ['Yes', 'No'])

cv_text = None  # Initialize cv_text to None

# File upload section allowing PDF and Word documents
uploaded_file = st.file_uploader("Choose a PDF or Word file", type=["pdf", "docx"])

if uploaded_file is not None:
    if uploaded_file.type == "application/pdf":
        cv_text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        cv_text = extract_text_from_docx(uploaded_file)

    if display_cv_text == 'Yes':
        st.write('Extracted text from your CV:')
        st.write(cv_text)

# Job URL input and display
job_url = st.text_input('Enter the URL of the job description')

if job_url:
    job_description = extract_text_from_url(job_url)

    if display_url_text == 'Yes':
        st.write('Extracted text from the job description:')
        st.write(job_description)

# Generate cover letter button
if st.button('Generate Cover Letter'):
    if uploaded_file is not None and job_url and api_key:
        cover_letter = generate_cover_letter(api_key, cv_text, job_description)
        st.write('Generated Cover Letter:')
        st.write(cover_letter)

        # PDF Download
        pdf_file = generate_pdf(cover_letter)
        st.download_button(
            label="Download Cover Letter as PDF",
            data=pdf_file,
            file_name="cover_letter.pdf",
            mime="application/pdf"
        )

        # DOCX Download
        docx_file = generate_docx(cover_letter)
        st.download_button(
            label="Download Cover Letter as DOCX",
            data=docx_file,
            file_name="cover_letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.write('Please upload a CV, enter a job description URL, and provide your OpenAI API Key.')
